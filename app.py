## Inicias descargado las librerias, al empquetar este script estas librerias ya no seran necesarias descargarce
import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import io

## Para darle mas estilo a la App.py
st.set_page_config(page_title='AutoReport', page_icon='smile')
# Modifica la definición de create_report
def create_report(template_file, data_dict, chart_data=None):
    st.write("Iniciando la creación del Informe.....")
    doc = Document(template_file)
    
    # Reemplaza las llaves en el documento
    for parrafo in doc.paragraphs:
        for key, value in data_dict.items():
            if f'{{{{{key}}}}}' in parrafo.text:
                st.write(f'Remplazando {key} con {value} en el informe')
                parrafo.text = parrafo.text.replace(f'{{{{{key}}}}}', str(value))
    
    # Generación del gráfico si chart_data no es None
    if chart_data is not None:
        st.write("Generando Gráfico....")
        plt.figure(figsize=(6, 4))
        plt.bar(chart_data['labels'], chart_data['values'])
        plt.title(chart_data['title'])
        plt.xlabel(chart_data.get('xlabel', ''))  # Proporcionar un valor por defecto
        plt.ylabel(chart_data.get('ylabel', ''))  # Proporcionar un valor por defecto

        # Guardar la figura en un buffer
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png')
        img_buffer.seek(0)
        st.write("Insertando Grafico en el documento.....")
        st.image(img_buffer)  # Mostrar el gráfico en Streamlit
        plt.close()  # Cerrar la figura para liberar memoria
        for parrafo in doc.paragraphs:
            for run in parrafo.runs:
                if '[Aquí se insertará el grafico]' in run.text:
                    run.text = run.text.replace('[Aquí se insertará el grafico]', '')
                    # Eliminar el texto de marcador
                    parrafo.text = parrafo.text.replace('[Aquí se insertará el gráfico]', '')
                    doc.add_picture(img_buffer, width=Inches(6))

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    st.write("Informe Creado con Éxito")
    return output


# Función principal para crear un título de la app web
def main():
    st.title("Generador de Informes desde Plantillas")

    template_file = st.file_uploader("Cargar Plantilla Word", type="docx")
    data_file = st.file_uploader("Cargar Datos", type=["xlsx", "csv"])
    
    df = None
    if template_file and data_file:
        st.success("Archivos Cargados Exitosamente")
        
        # Procesar el archivo solo si se ha cargado
        if data_file.name.endswith(".csv"):
            df = pd.read_csv(data_file)
        else:
            df = pd.read_excel(data_file)

        # Agregar un subtítulo
        st.subheader("Datos Cargados")
        
        # Muestra el DataFrame
        st.dataframe(df)  # Mostrar tabla para analizar

    if df is not None:
        # Selecciona la fila para gestionar el Informe
        row_index = st.selectbox("Selecciona la fila para gestionar el Informe", options=range(len(df)))
        data_selecionada = df.iloc[row_index].to_dict()  # Guarda la fila como diccionario
        
        generate_chart = st.checkbox("Generar Grafico")
        chart_data = None
        
        if generate_chart:
            chart_title = st.text_input("Titulo del Grafico", "Grafico de Datos")
            x_column = st.selectbox("Columna para Eje X", options= df.columns)
            y_column = st.selectbox("Columna para Eje Y", options= df.columns)
            
            chart_data = {
                'title': chart_title,
                'labels': df[x_column].tolist(),
                'values': df[y_column].tolist(),
                'xlabel': x_column,
                'ylabel': y_column
            }
            
            st.write("Datos del Grafico:", chart_data)
        
        if st.button("Generar Informe"):
            output = create_report(template_file, data_selecionada, chart_data)  # Pasa el diccionario a la función
            st.download_button("Descargar Informe", output, "Informe_Generado.docx",
                               "application/vnd.openxmlformants-officiedocument.wordprocessingml.document")
            
# Comprobación de si el script se ejecuta como principal
if __name__ == "__main__":
    main()